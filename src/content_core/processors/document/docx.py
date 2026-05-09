import asyncio

from docx import Document  # type: ignore

from content_core.logging import logger


async def extract_docx_content_detailed(file_path):
    """Extract content from DOCX file"""

    def _extract():
        try:
            doc = Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue

                style = paragraph.style.name if paragraph.style else "Normal"
                text = paragraph.text.strip()

                # Get paragraph formatting
                p_format = paragraph.paragraph_format
                indent = p_format.left_indent or 0

                # Convert indent to spaces (1 level = 4 spaces)
                indent_level = 0
                if hasattr(indent, "pt"):
                    indent_level = int(indent.pt / 72)  # 72 points = 1 inch
                indent_spaces = " " * (indent_level * 4)

                # Handle different types of formatting
                if "Heading" in style:
                    level = style[-1] if style[-1].isdigit() else "1"
                    heading_marks = "#" * int(level)
                    content.append(f"\n{heading_marks} {text}\n")

                # Handle bullet points
                elif (
                    paragraph.style
                    and hasattr(paragraph.style, "name")
                    and paragraph.style.name.startswith("List")
                ):
                    # Numbered list
                    if (
                        hasattr(paragraph._p, "pPr")
                        and paragraph._p.pPr is not None
                        and hasattr(paragraph._p.pPr, "numPr")
                        and paragraph._p.pPr.numPr is not None
                    ):
                        # Try to get the actual number
                        try:
                            if (
                                hasattr(paragraph._p.pPr.numPr, "numId")
                                and paragraph._p.pPr.numPr.numId is not None
                                and hasattr(paragraph._p.pPr.numPr.numId, "val")
                            ):
                                number = paragraph._p.pPr.numPr.numId.val
                                content.append(f"{indent_spaces}{number}. {text}")
                            else:
                                content.append(f"{indent_spaces}1. {text}")
                        except Exception:
                            content.append(f"{indent_spaces}1. {text}")
                    # Bullet list
                    else:
                        content.append(f"{indent_spaces}* {text}")

                else:
                    # Handle text formatting
                    formatted_text = []
                    for run in paragraph.runs:
                        if run.bold:
                            formatted_text.append(f"**{run.text}**")
                        elif run.italic:
                            formatted_text.append(f"*{run.text}*")
                        else:
                            formatted_text.append(run.text)

                    content.append(f"{indent_spaces}{''.join(formatted_text)}")

            return "\n\n".join(content)

        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
            return None

    return await asyncio.get_event_loop().run_in_executor(None, _extract)


async def get_docx_info(file_path):
    """Get DOCX metadata and content"""

    async def _get_info():
        try:
            doc = Document(file_path)

            # Extract core properties if available
            core_props = {
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "title": doc.core_properties.title,
                "subject": doc.core_properties.subject,
                "keywords": doc.core_properties.keywords,
                "category": doc.core_properties.category,
                "comments": doc.core_properties.comments,
            }

            # Get document content
            content = await extract_docx_content_detailed(file_path)

            # Get document statistics
            stats = {
                "paragraph_count": len(doc.paragraphs),
                "word_count": sum(
                    len(p.text.split()) for p in doc.paragraphs if p.text.strip()
                ),
                "character_count": sum(
                    len(p.text) for p in doc.paragraphs if p.text.strip()
                ),
            }

            return {"metadata": core_props, "content": content, "statistics": stats}

        except Exception as e:
            logger.error(f"Failed to get DOCX info: {e}")
            return None

    return await _get_info()
