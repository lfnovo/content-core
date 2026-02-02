"""DOCX extraction engines and quality scoring."""

import asyncio
import re
from typing import Any, Dict, List, Optional

from ..base import ContentAnalyzer, Engine, QualityScore, QualityScorer
from ..test_data.docx_expected import BENCHMARK_DOCX_EXPECTED


# --- Quality Scorer ---


class DOCXQualityScorer(QualityScorer):
    """Score DOCX extraction quality against expected content."""

    file_type = "docx"

    def score(self, content: str, file_name: str) -> Optional[QualityScore]:
        """Score extraction quality against expected benchmark.docx content."""
        if file_name != "benchmark.docx":
            return None

        content_lower = content.lower()
        details = {}

        # Check title
        details["title"] = BENCHMARK_DOCX_EXPECTED["title"].lower() in content_lower

        # Check sections
        for section in BENCHMARK_DOCX_EXPECTED["sections"]:
            key = f"section_{section.lower().replace(' ', '_')}"
            details[key] = section.lower() in content_lower

        # Check subsections
        for subsection in BENCHMARK_DOCX_EXPECTED["subsections"]:
            key = f"subsection_{subsection.lower().replace(' ', '_')}"
            details[key] = subsection.lower() in content_lower

        # Check formulas (case-sensitive for math)
        for formula_name, patterns in BENCHMARK_DOCX_EXPECTED["formulas"]:
            key = f"formula_{formula_name}"
            details[key] = any(pattern in content for pattern in patterns)

        # Check table data
        for data in BENCHMARK_DOCX_EXPECTED["table_data"]:
            key = f"table_{data.lower().replace(' ', '_')}"
            details[key] = data.lower() in content_lower

        # Check figure references
        for fig in BENCHMARK_DOCX_EXPECTED["figure"]:
            key = f"figure_{fig.lower().replace(' ', '_')}"
            details[key] = fig.lower() in content_lower

        # Check DOCX-specific elements
        for fmt in BENCHMARK_DOCX_EXPECTED.get("formatting", []):
            key = f"formatting_{fmt.lower()}"
            details[key] = fmt.lower() in content_lower

        for item in BENCHMARK_DOCX_EXPECTED.get("lists", []):
            key = f"list_{item.lower().replace(' ', '_')}"
            details[key] = item.lower() in content_lower

        found = sum(1 for v in details.values() if v)
        total = len(details)
        score = found / total if total > 0 else 0.0

        return QualityScore(score=score, found=found, total=total, details=details)


# --- Content Analyzer ---


class DOCXContentAnalyzer(ContentAnalyzer):
    """Analyze DOCX extracted content for structure metrics."""

    file_type = "docx"

    def analyze(self, content: str) -> Dict[str, Any]:
        """Extract metrics from content."""
        return {
            "headers_count": self._count_headers(content),
            "tables_count": self._count_tables(content),
            "formulas_count": self._count_formulas(content),
            "lists_count": self._count_lists(content),
        }

    def _count_headers(self, content: str) -> int:
        """Count markdown headers (## style)."""
        return len(re.findall(r"^#{1,6}\s+.+$", content, re.MULTILINE))

    def _count_tables(self, content: str) -> int:
        """Count markdown tables (lines starting with |)."""
        table_rows = re.findall(r"^\|.+\|$", content, re.MULTILINE)
        if not table_rows:
            return 0
        separators = len(re.findall(r"^\|[\s\-:|]+\|$", content, re.MULTILINE))
        return max(separators, 1) if table_rows else 0

    def _count_formulas(self, content: str) -> int:
        """Count LaTeX formulas ($$ or $ delimited)."""
        block_formulas = len(re.findall(r"\$\$[^$]+\$\$", content))
        inline_formulas = len(re.findall(r"(?<!\$)\$(?!\$)[^$]+\$(?!\$)", content))
        return block_formulas + inline_formulas

    def _count_lists(self, content: str) -> int:
        """Count list items (- or * or numbered)."""
        bullet_items = len(re.findall(r"^[\s]*[-*]\s+.+$", content, re.MULTILINE))
        numbered_items = len(re.findall(r"^[\s]*\d+\.\s+.+$", content, re.MULTILINE))
        return bullet_items + numbered_items


# --- Engines ---


class PythonDocxEngine(Engine):
    """Basic python-docx extraction - extracts text without markdown formatting.

    Fast and lightweight, but doesn't preserve formatting or structure well.
    """

    name = "python-docx"
    supported_types = ["docx"]

    async def extract(self, file_path: str, options: Dict[str, Any]) -> str:
        """Extract DOCX using python-docx."""
        from docx import Document

        def _extract():
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Try to detect headings by style
                    if para.style and para.style.name.startswith("Heading"):
                        level = 1
                        try:
                            level = int(para.style.name.split()[-1])
                        except (ValueError, IndexError):
                            pass
                        paragraphs.append(f"{'#' * level} {text}")
                    else:
                        paragraphs.append(text)

            # Extract tables
            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append("| " + " | ".join(cells) + " |")
                    if len(table_lines) == 1:
                        # Add separator after header
                        table_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
                if table_lines:
                    paragraphs.append("\n".join(table_lines))

            return "\n\n".join(paragraphs)

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract)


class DoclingDocxEngine(Engine):
    """Docling extraction - best quality for DOCX files.

    Uses docling's document understanding pipeline with OCR and table detection.
    """

    name = "docling"
    supported_types = ["docx"]

    async def extract(self, file_path: str, options: Dict[str, Any]) -> str:
        """Extract DOCX using docling."""
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()

        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, converter.convert, file_path)

        doc = result.document
        return doc.export_to_markdown()


# --- Factory Functions ---


def get_docx_engines(names: Optional[List[str]] = None) -> List[Engine]:
    """Get DOCX extraction engines by name.

    Args:
        names: List of engine names to get. If None, returns all available engines.

    Returns:
        List of Engine instances
    """
    all_engines = {
        "python-docx": PythonDocxEngine(),
        "docling": DoclingDocxEngine(),
    }

    if names is None:
        return list(all_engines.values())

    engines = []
    for name in names:
        if name in all_engines:
            engines.append(all_engines[name])
        else:
            raise ValueError(f"Unknown DOCX engine: {name}. Available: {list(all_engines.keys())}")
    return engines


# List of available DOCX engine names
AVAILABLE_DOCX_ENGINES = ["python-docx", "docling"]
